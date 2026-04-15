"""
Convert PROJECT_DOCUMENTATION.md to a professional Word document.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path
import re

doc = Document()

# -- Page margins --
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# Helper functions
def add_heading(text, level=1):
    h = doc.add_heading(text, level=level)
    return h

def add_para(text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    return p

def add_table(headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header row
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for p in hdr_cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.size = Pt(10)
    # Data rows
    for row_data in rows:
        row_cells = table.add_row().cells
        for i, val in enumerate(row_data):
            row_cells[i].text = str(val)
            for p in row_cells[i].paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
    doc.add_paragraph()
    return table

# ============================================================
# TITLE PAGE
# ============================================================
for _ in range(6):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SMART ENVIRONMENT ENERGY OPTIMIZER")
run.bold = True
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0, 51, 102)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("Using Deep Learning-Based Occupancy Estimation\nand Reinforcement Learning Control")
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(51, 51, 51)

doc.add_paragraph()
doc.add_paragraph()

info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = info.add_run("Complete Project Documentation\nDeep Learning + Reinforcement Learning\n\n[Team Member Names]\n[College Name]\n[Date]")
run.font.size = Pt(12)
run.font.color.rgb = RGBColor(80, 80, 80)

doc.add_page_break()

# ============================================================
# TABLE OF CONTENTS
# ============================================================
add_heading("Table of Contents", 1)
toc_items = [
    "1. Project Overview",
    "2. System Architecture",
    "3. Deep Learning Component — Occupancy Estimation",
    "4. Reinforcement Learning Component — Energy Optimization",
    "5. Integration & Demo Pipeline",
    "6. Results & Evaluation",
    "7. Why We Chose What We Chose — All The Whys",
    "8. Known Limitations & Honest Assessment",
    "9. How to Run",
    "10. File Structure",
    "11. References",
]
for item in toc_items:
    doc.add_paragraph(item, style='List Number')
doc.add_page_break()

# ============================================================
# 1. PROJECT OVERVIEW
# ============================================================
add_heading("1. Project Overview", 1)

add_heading("The Problem (In Simple Words)", 2)
add_para("Imagine your college building. Every day, lights, fans, and ACs run at full power in EVERY room — even when classrooms are empty. This wastes a massive amount of electricity and money.")
add_para("Real-world example: A lecture hall with 8 tube lights (40W each), 2 fans (75W each), and 1 AC (1500W) consumes 1,970 watts. If it runs for 10 hours daily, that's 19.7 kWh per day — just ONE room. With 5 rooms, that's ~100 kWh/day or Rs 800/day wasted when rooms are empty.", italic=True)

add_heading("Our Solution", 2)
add_para("We built an AI-powered energy management system that:")
doc.add_paragraph("SEES how many people are in each room (using a camera + Deep Learning)", style='List Bullet')
doc.add_paragraph("DECIDES what to turn on/off (using a trained RL agent)", style='List Bullet')
doc.add_paragraph("SAVES energy while keeping everyone comfortable", style='List Bullet')

add_para("")
add_para('Layman Analogy: Think of it like a smart thermostat, but instead of just temperature, our AI watches the room through a camera, counts people, and then determines: "There are only 5 students in Room 102 — I\'ll keep lights at half and turn AC off. But Room 103 has 35 students — everything goes to full power."', italic=True)

add_heading("Key Achievements", 2)
add_table(
    ["Metric", "Value"],
    [
        ["Energy Savings", "69-76% compared to always-on baseline"],
        ["DL Accuracy (MAE)", "0.69 (less than 1 person error on average)"],
        ["DL Accuracy (MAPE)", "2.3% error"],
        ["Cost Savings", "Rs 540-665 per day"],
    ]
)

doc.add_page_break()

# ============================================================
# 2. SYSTEM ARCHITECTURE
# ============================================================
add_heading("2. System Architecture", 1)

add_heading("High-Level Flow", 2)
add_para("Camera/Video → DL Pipeline → Occupancy Count → RL Agent → Device Control → Dashboard")
add_para("")
add_para("The system takes a live video feed, processes each frame through a deep learning pipeline to count people, feeds the occupancy data to a reinforcement learning agent that decides optimal device settings, and displays everything on a real-time dashboard.")

add_heading("Component Breakdown", 2)
add_table(
    ["Component", "Technology", "Purpose"],
    [
        ["Feature Extractor", "YOLOv8-small (frozen)", "Extract visual features from camera frames"],
        ["Density Head", "3-layer Dilated CNN", "Generate density maps → count people"],
        ["RL Agent", "PPO (Stable-Baselines3)", "Decide device levels for each room"],
        ["Environment", "Custom Gymnasium Env", "Simulate college floor with 5 rooms"],
        ["Dashboard", "Matplotlib (TkAgg)", "Real-time visualization"],
    ]
)

add_heading("Data Flow (Step by Step)", 2)
steps = [
    "Camera captures a frame (640×480 pixels)",
    "YOLOv8 backbone extracts features → [B, 128, 60, 80] tensor",
    "Density Head produces a density map → sum gives people count",
    "Auto-calibration removes scene-specific background noise",
    "Count is normalized to occupancy percentage (0-100%)",
    "RL agent observes occupancy + device states (27-dim vector)",
    "Agent outputs device levels (OFF / HALF / FULL) for each device",
    "Comfort override ensures minimum comfort when people are present",
    "Dashboard displays everything in real-time",
]
for i, step in enumerate(steps, 1):
    doc.add_paragraph(f"{i}. {step}")

doc.add_page_break()

# ============================================================
# 3. DEEP LEARNING COMPONENT
# ============================================================
add_heading("3. Deep Learning Component — Occupancy Estimation", 1)
add_para("This section covers the deep learning pipeline for counting people in indoor rooms.", italic=True)

add_heading("3.1 Problem Statement", 2)
add_para("Estimate the number of people in an indoor room from a single camera frame, without requiring person detection or tracking.")

add_heading("3.2 Approach: Density Map Regression", 2)
add_para("Instead of detecting individual people (which fails in crowded scenes due to occlusion), we use density map regression:")
doc.add_paragraph("For each annotated head position, place a Gaussian blob on a map", style='List Bullet')
doc.add_paragraph("Train a CNN to output a density map from an image", style='List Bullet')
doc.add_paragraph("Sum of all pixels in the density map = person count", style='List Bullet')

add_para("")
add_para("Why density maps instead of direct counting?", bold=True)
doc.add_paragraph("Direct regression ('this image has 23 people') doesn't learn spatial information", style='List Bullet')
doc.add_paragraph("Density maps preserve WHERE people are, making the model much more robust", style='List Bullet')
doc.add_paragraph("The model can generalize to different crowd densities", style='List Bullet')

add_para("")
add_para('Layman analogy: Instead of asking "how many grains of rice are in this bowl?", we first create a heat map showing where rice is concentrated. Adding up the heat map gives us the total count. This is more reliable than trying to count individual grains.', italic=True)

add_heading("3.3 Architecture", 2)

add_para("Backbone: YOLOv8-small (FROZEN)", bold=True)
add_para("Input Image [3, 480, 640] → YOLOv8-small backbone (pretrained on COCO, weights FROZEN) → Feature Map [128, 60, 80]")
doc.add_paragraph("We use YOLOv8-small as a feature extractor only", style='List Bullet')
doc.add_paragraph("Pretrained on COCO (Common Objects in Context) — 330K images", style='List Bullet')
doc.add_paragraph("All backbone weights are FROZEN — we don't retrain them", style='List Bullet')
doc.add_paragraph("Why freeze? The backbone already knows how to extract visual features. Training it would require 100x more data.", style='List Bullet')
doc.add_paragraph("We extract from layer 4 (C2f block, stride-8 downsampling) → 128 channels at 1/8 resolution", style='List Bullet')

add_para("")
add_para("Density Head (TRAINABLE)", bold=True)
add_para("Layer 1: Conv2d(128→64, kernel=3×3, dilation=1, padding=1) + ReLU")
add_para("Layer 2: Conv2d(64→32, kernel=3×3, dilation=2, padding=2) + ReLU")
add_para("Layer 3: Conv2d(32→1, kernel=3×3, dilation=2, padding=2) + ReLU")
add_para("Output: Bilinear Upsample to [1, 480, 640] → Density Map")

add_para("")
add_para("Why dilated convolutions?", bold=True)
add_para("Dilated (atrous) convolutions increase the receptive field without adding parameters. With dilation=2, a 3×3 kernel effectively covers a 5×5 area. This helps the model understand context around each person without being computationally expensive.")

add_para("Why ReLU at the end (not sigmoid)?", bold=True)
add_para("Density values should be non-negative (can't have -0.5 people) but not bounded above. Sigmoid would cap at 1.0, which is wrong for density values.")

add_heading("3.4 Dataset: Mall Dataset", 2)
add_table(
    ["Property", "Value"],
    [
        ["Source", "CUHK (Chinese University of Hong Kong)"],
        ["Total Frames", "2000 video frames from mall surveillance"],
        ["Resolution", "640 × 480 pixels"],
        ["People per frame", "13 to 53 (mean: 31.2)"],
        ["Annotations", "Head (x,y) positions for every person"],
        ["Scene", "Indoor — overhead surveillance view"],
        ["Total annotations", "60,000+ head positions"],
    ]
)

add_para("Why Mall Dataset and not ShanghaiTech?", bold=True)
add_para("We initially trained on ShanghaiTech Part B (outdoor crowd scenes — streets, plazas, stadiums). When applied to indoor rooms, the model hallucinated 30-40 people in empty classrooms because outdoor textures are completely different from indoor environments.")
add_para("The Mall Dataset is indoor surveillance — very similar to our target scenario (classroom cameras). After retraining, MAE improved from 19.22 to 0.69.")

add_heading("3.5 Training Details", 2)

add_para("Preprocessing: Density Map Generation", bold=True)
add_para("For each frame with N annotated head positions:")
doc.add_paragraph("Create a blank image (480×480 zeros)", style='List Number')
doc.add_paragraph("Place a value of 1.0 at each head (x, y) coordinate", style='List Number')
doc.add_paragraph("Apply a Gaussian filter (σ=4.0, kernel_size=15)", style='List Number')
doc.add_paragraph("Result: density map where sum of all pixels ≈ N", style='List Number')

add_para("")
add_para("Training Hyperparameters", bold=True)
add_table(
    ["Hyperparameter", "Value", "Justification"],
    [
        ["Epochs", "50 (max)", "Sufficient for 1600 samples; early stopping prevents overfitting"],
        ["Batch Size", "4", "Small for CPU training; high-res maps use memory; adds regularization"],
        ["Learning Rate", "1e-4", "Standard for Adam fine-tuning; not too aggressive"],
        ["Optimizer", "Adam", "Adaptive per-parameter rates; works well with small datasets"],
        ["Early Stopping Patience", "12 epochs", "Stop if no improvement for 12 consecutive epochs"],
        ["Train/Val Split", "80/20", "1600 train + 400 validation frames"],
        ["Augmentation", "Horizontal flip + color jitter", "Effectively doubles data; handles lighting changes"],
        ["Seed", "42", "Convention for reproducibility"],
        ["Gaussian σ", "4.0", "Each person affects ~15×15 pixels; standard for medium-density"],
    ]
)

add_para("Loss Function: DensityLoss", bold=True)
add_para("Loss = MSE(predicted_density, ground_truth_density) + 0.01 × |predicted_count - actual_count|")
add_para("")
doc.add_paragraph("MSE (pixel-level): Forces the density map shape to match ground truth", style='List Bullet')
doc.add_paragraph("Counting loss (weight=0.01): Directly penalizes wrong total counts", style='List Bullet')
doc.add_paragraph("Why both? MSE alone might produce beautiful maps that sum to wrong numbers. Count loss ensures correctness.", style='List Bullet')

add_para("")
add_para("Learning Rate Scheduler", bold=True)
add_para("ReduceLROnPlateau: If validation loss doesn't improve for 5 epochs, halve the learning rate. Factor=0.5, Patience=5 epochs. This automatically fine-tunes as training progresses.")

add_heading("3.6 Auto-Calibration (Noise Removal)", 2)
add_para("The density model has inherent 'background noise' — it sees texture patterns in walls and furniture and assigns small density values. This varies per scene.")
add_para("")
add_para("Our solution: Auto-calibration using the first 5 frames:", bold=True)
doc.add_paragraph("Run the model on the first 5 frames of any video", style='List Number')
doc.add_paragraph("Take the minimum raw count × 0.9 = noise floor", style='List Number')
doc.add_paragraph("All subsequent predictions subtract this noise floor", style='List Number')

add_para("")
add_para("Example:", bold=True)
add_table(
    ["Video", "Raw Count", "Auto Noise Floor", "Cleaned Count", "Result"],
    [
        ["empty.mp4 (0 people)", "29", "25.2", "~4", "Close to 0 ✓"],
        ["aarya.mp4 (7 people)", "18", "~15", "~8", "Close to 7 ✓"],
        ["hackx.mp4 (15 people)", "30", "~18", "~15", "Accurate ✓"],
    ]
)

add_heading("3.7 Evaluation Results", 2)
add_para("Quantitative Results on Mall Dataset (2000 frames):", bold=True)
add_table(
    ["Metric", "Value", "Meaning"],
    [
        ["MAE", "0.69", "Model is off by less than 1 person on average"],
        ["RMSE", "1.07", "Root mean squared error (penalizes large errors)"],
        ["MAPE", "2.3%", "Only 2.3% percentage error — 97.7% accurate"],
        ["GT Range", "13 - 53", "Ground truth people per frame"],
        ["Pred Range", "14.0 - 53.6", "Model's predicted range"],
    ]
)

doc.add_page_break()

# ============================================================
# 4. REINFORCEMENT LEARNING COMPONENT
# ============================================================
add_heading("4. Reinforcement Learning Component — Energy Optimization", 1)
add_para("This section covers the RL agent that learns to control devices optimally.", italic=True)

add_heading("4.1 Problem Formulation", 2)
add_para("Goal: Learn a policy that controls lights, fans, and ACs across 5 rooms to minimize energy consumption while maintaining occupant comfort.", bold=True)
add_para("")
add_para('Layman Analogy: Imagine you\'re a building manager. Every 6 minutes, you walk through the building, check how many people are in each room, and decide: should the lights be off, at half, or full? Same for fans and AC. You want to save electricity but not make students uncomfortable. Over time, you learn the patterns — "Room 101 is always empty at 3 PM, so I can turn everything off." Our RL agent learns this same intuition, but automatically.', italic=True)

add_heading("4.2 Environment Design (CollegeFloorEnv)", 2)

add_para("Room Configuration", bold=True)
add_table(
    ["Room", "Type", "Capacity", "Devices", "Max Power"],
    [
        ["101", "Lecture Hall", "60", "8 lights, 2 fans, 1 AC", "1,970W"],
        ["102", "Lecture Hall", "60", "8 lights, 2 fans, 1 AC", "1,970W"],
        ["103", "Computer Lab", "40", "6 lights, 1 AC, 40 PCs", "4,740W"],
        ["104", "Tutorial Room", "30", "4 lights, 2 fans", "310W"],
        ["105", "Staff Room", "10", "4 lights, 1 AC", "1,660W"],
        ["Corridor", "—", "—", "4 lights (always 50%)", "80W"],
    ]
)
add_para("Total always-on baseline: 108.1 kWh/day (if everything runs at full for 10 hours)")

add_heading("MDP Formulation", 3)
add_para("State Space (Observation): Box(27)", bold=True)
doc.add_paragraph("Per room (5 rooms × 5 values = 25): [occupancy, activity, light_level, fan_level, ac_level]", style='List Bullet')
doc.add_paragraph("Global (2 values): [outside_temperature_normalized, time_of_day_normalized]", style='List Bullet')
doc.add_paragraph("All values normalized to [0, 1]", style='List Bullet')

add_para("")
add_para("Action Space: MultiDiscrete([3] × 12)", bold=True)
doc.add_paragraph("12 controllable devices (each room's available devices)", style='List Bullet')
doc.add_paragraph("Each device has 3 levels: 0 (OFF), 1 (HALF), 2 (FULL)", style='List Bullet')
doc.add_paragraph("Total possible actions: 3^12 = 531,441 combinations", style='List Bullet')

add_para("")
add_para("Episode: 96 timesteps = 10 hours (8 AM to 6 PM), each step ≈ 6.25 minutes")

add_heading("Occupancy Simulation", 3)
add_para("Occupancy follows realistic Gaussian curve patterns:")
doc.add_paragraph("Room 101 (Lecture Hall): Peaks at 9-11 AM and 1-3 PM (class hours)", style='List Bullet')
doc.add_paragraph("Room 103 (Computer Lab): Peaks at 2-5 PM (lab sessions)", style='List Bullet')
doc.add_paragraph("Room 104 (Tutorial): Short scattered peaks throughout the day", style='List Bullet')
doc.add_paragraph("Room 105 (Staff Room): Steady low occupancy 9 AM - 5 PM", style='List Bullet')
add_para("Each episode adds a random phase offset (±5%) so the agent doesn't memorize exact patterns. Gaussian noise (σ=0.1) is also added for stochasticity.")

add_heading("4.3 Reward Function", 2)
add_para("Reward = α × Comfort − β × Energy_Cost − γ × Adjustment_Penalty", bold=True)
add_table(
    ["Term", "Weight", "Purpose"],
    [
        ["α × Comfort", "α = 1.0", "Positive reward for keeping occupants comfortable"],
        ["β × Energy_Cost", "β = 0.6", "Penalty for consuming energy (Rs 8/kWh)"],
        ["γ × Adjustment", "γ = 0.1", "Penalty for frequently switching devices"],
    ]
)

add_para("How Comfort is Calculated:", bold=True)
add_para("Comfort = 1 − mean(max(required_level − actual_level, 0)) across devices")
add_table(
    ["Occupancy", "Required Level", "Example"],
    [
        ["> 70%", "1.0 (FULL)", "40 students → everything at max"],
        ["30-70%", "0.5 (HALF)", "15 students → devices at half"],
        ["< 30%", "0.0 (OFF OK)", "Empty → agent can turn everything off"],
    ]
)
add_para("Important: Only UNDER-provision hurts comfort. Over-provision (lights on in empty room) doesn't hurt comfort — it only wastes energy. This teaches the agent to save energy proactively.", italic=True)

add_heading("Why β = 0.6? (Beta Sweep Experiment)", 3)
add_para("We tested three beta values:", bold=True)
add_table(
    ["β (Beta)", "Energy (kWh/day)", "Comfort", "Behavior"],
    [
        ["0.3", "26.4", "84.0%", "Comfort-focused, wastes some energy"],
        ["0.6 (CHOSEN)", "25.1", "81.2%", "Best balance of savings vs comfort"],
        ["0.9", "26.5", "83.3%", "Too aggressive, unstable behavior"],
    ]
)

add_heading("4.4 PPO Algorithm", 2)

add_para("Why PPO?", bold=True)
add_table(
    ["Algorithm", "Pros", "Cons", "Verdict"],
    [
        ["DQN", "Simple", "Can't handle MultiDiscrete (531K actions)", "✗"],
        ["A2C", "Fast", "High variance, unstable", "✗"],
        ["PPO", "Stable, handles MultiDiscrete, sample efficient", "Slower", "✓ CHOSEN"],
        ["SAC", "State-of-the-art for continuous", "Requires continuous actions", "✗"],
    ]
)

add_para("PPO (Proximal Policy Optimization) Key Properties:", bold=True)
doc.add_paragraph("Clipped objective prevents catastrophic policy updates", style='List Bullet')
doc.add_paragraph("Works with MultiDiscrete action spaces naturally", style='List Bullet')
doc.add_paragraph("Industry standard for real-world control problems (robotics, games, energy)", style='List Bullet')

add_para("")
add_para("PPO Hyperparameters", bold=True)
add_table(
    ["Parameter", "Value", "Justification"],
    [
        ["Total timesteps", "200,000", "≈2,083 episodes. Sufficient for convergence"],
        ["Learning rate", "3e-4", "SB3 default, well-tested for PPO"],
        ["n_steps", "2048", "~21 episodes per update. Stable gradients"],
        ["batch_size", "64", "Mini-batch size. Power of 2"],
        ["n_epochs", "10", "Passes over rollout buffer per update"],
        ["Policy", "MlpPolicy (64×64)", "2-layer MLP. Tabular data, not images"],
        ["Discount (γ)", "0.99", "Agent cares about long-term savings"],
        ["Seed", "42", "Reproducibility"],
    ]
)

add_para('Layman analogy for training: The agent plays 21 full "games" (episodes of 96 steps each), remembers what happened, then reviews all 2048 steps to improve its strategy. It repeats this ~100 times (200K total steps) until it becomes expert at energy management.', italic=True)

add_heading("4.5 Comfort Override (Safety Layer)", 2)
add_para("The PPO agent sometimes turns everything off aggressively. In reality, students can't sit in the dark! We added a safety override:", bold=True)
add_table(
    ["Occupancy", "Override Action"],
    [
        ["> 60% (crowded)", "All devices → FULL"],
        ["25-60% (moderate)", "All devices → at least HALF"],
        ["5-25% (few people)", "Lights + Fan → HALF, AC stays off"],
        ["< 5% (empty)", "Agent's choice (may turn everything off)"],
    ]
)

add_heading("4.6 RL Training Results", 2)
add_table(
    ["Metric", "Always-On Baseline", "PPO Agent", "Improvement"],
    [
        ["Energy/day", "108.1 kWh", "25.1 kWh", "76.8% savings"],
        ["Cost/day", "Rs 864.80", "Rs 200.74", "Rs 664 saved"],
        ["Comfort", "100%", "81.2%", "Acceptable trade-off"],
    ]
)

doc.add_page_break()

# ============================================================
# 5. INTEGRATION
# ============================================================
add_heading("5. Integration & Demo Pipeline", 1)
add_para("The demo.py script ties everything together:")
doc.add_paragraph("Loads the trained PPO model from ppo_college_floor.zip", style='List Number')
doc.add_paragraph("Initializes the CollegeFloorEnv gymnasium environment", style='List Number')
doc.add_paragraph("Opens the video source (file or webcam) for Room 102", style='List Number')
doc.add_paragraph("Runs the DL pipeline on each frame to get occupancy", style='List Number')
doc.add_paragraph("Feeds occupancy to the RL agent for device decisions", style='List Number')
doc.add_paragraph("Applies comfort override for human safety", style='List Number')
doc.add_paragraph("Updates the real-time dashboard", style='List Number')
doc.add_paragraph("Loops until video ends or user presses Ctrl+C", style='List Number')

add_para("")
add_para("Dashboard Features:", bold=True)
doc.add_paragraph("5 room cards showing people count, occupancy %, device status (L/F/AC), power draw", style='List Bullet')
doc.add_paragraph("Live metrics panel: power, energy, cost, savings %, comfort %", style='List Bullet')
doc.add_paragraph("Occupancy bar chart comparing all rooms", style='List Bullet')
doc.add_paragraph("Savings summary showing baseline vs AI-optimized energy", style='List Bullet')

doc.add_page_break()

# ============================================================
# 6. RESULTS
# ============================================================
add_heading("6. Results & Evaluation", 1)

add_heading("6.1 DL Model Evaluation", 2)
add_table(
    ["Metric", "Value", "Meaning"],
    [
        ["MAE", "0.69", "Less than 1 person error on average"],
        ["RMSE", "1.07", "Root mean squared error"],
        ["MAPE", "2.3%", "97.7% accurate"],
    ]
)
add_para("Evaluation plots generated: evaluation_results/evaluation_plots.png (scatter, error distribution, timeline, residuals) and evaluation_results/sample_predictions.png (visual predictions on 8 frames)")

add_heading("6.2 RL Agent Performance", 2)
add_table(
    ["Scenario", "Energy (kWh)", "Cost (Rs)", "Savings"],
    [
        ["Always-On Baseline", "108.1", "864.80", "—"],
        ["PPO Agent (simulated)", "25.1", "200.74", "76.8%"],
        ["PPO + Video (aarya.mp4)", "40.7", "325.33", "62.4%"],
        ["PPO + Video (hackx.mp4)", "32.9", "262.92", "69.6%"],
    ]
)

doc.add_page_break()

# ============================================================
# 7. WHY WE CHOSE WHAT WE CHOSE
# ============================================================
add_heading("7. Why We Chose What We Chose — All The Whys", 1)

whys = [
    ("Why seed = 42?", "42 is the most commonly used random seed in ML research. It's a cultural convention from 'The Hitchhiker's Guide to the Galaxy' where 42 is the 'Answer to the Ultimate Question of Life'. Any fixed integer works — the point is reproducibility: same code = same results."),
    ("Why YOLOv8 and not ResNet/VGG?", "YOLOv8-s has 11.2M parameters (small). VGG-16 has 138M. YOLOv8 is pretrained on COCO (80 classes including 'person') and runs at ~50ms/frame on CPU. It's the latest version with modern C2f blocks."),
    ("Why freeze the backbone?", "We only have 2000 training images. Training 11.2M parameters on 2000 images = massive overfitting. The backbone already extracts excellent features from COCO pretraining. We only train the density head (5,793 parameters)."),
    ("Why 50 epochs?", "With 1600 samples and batch_size=4, each epoch = 400 iterations. Early stopping with patience=12 typically stops at epoch 25-35. Beyond 50 risks overfitting."),
    ("Why batch size 4?", "CPU training — larger batches are slow. Small batches add implicit regularization via noisier gradients. Density maps (60×80) consume significant memory."),
    ("Why learning rate 1e-4?", "Standard for Adam fine-tuning. 1e-3 caused instability. 1e-5 was too slow. ReduceLROnPlateau automatically halves it when progress stalls."),
    ("Why Gaussian σ=4.0?", "Each person affects ~15×15 pixels (3σ rule). Too small (σ=1) = sparse maps. Too large (σ=15) = blurred people. Literature standard is σ=3-5."),
    ("Why PPO and not Q-learning?", "Our action space is 3^12 = 531,441 possible actions. Q-learning needs one Q-value per action — impractical. PPO's policy network scales naturally."),
    ("Why 200,000 timesteps?", "≈2,083 episodes, ~97 policy updates. Convergence typically occurs around 50 updates. 200K is the standard starting point for PPO."),
    ("Why Rs 8/kWh?", "Average Indian commercial electricity rate is Rs 6-10/kWh. Rs 8 is a representative mid-range value."),
    ("Why α=1.0, β=0.6, γ=0.1?", "α=1.0: Full weight on comfort. β=0.6: Found via beta sweep as the optimal balance. γ=0.1: Small penalty for switching to prevent flickering."),
    ("Why 96 timesteps per episode?", "10-hour workday (8 AM - 6 PM) ÷ 96 = 6.25 min/step. Matches real building management system intervals (5-15 min)."),
]

for q, a in whys:
    add_para(q, bold=True)
    add_para(a)
    add_para("")

doc.add_page_break()

# ============================================================
# 8. LIMITATIONS
# ============================================================
add_heading("8. Known Limitations & Honest Assessment", 1)

add_heading("8.1 DL Limitations", 2)
limitations_dl = [
    ("Domain Gap", "Model trained on one mall camera. Different scenes have different noise levels. Mitigation: Auto-calibration subtracts scene-specific noise."),
    ("Empty Room Issue", "Empty rooms still show raw count of ~25-30. After calibration, ~3-5 (not zero). Fundamental limitation of density-based counting."),
    ("Single Camera Only", "Currently one camera → one room (Room 102). Other rooms simulated. 5 cameras would need 5× inference time."),
    ("No Person Re-identification", "Model counts density, not individuals. Can't tell if Person A left and B entered. For energy management, we only need count — acceptable."),
]
for title, desc in limitations_dl:
    add_para(title, bold=True)
    add_para(desc)

add_heading("8.2 RL Limitations", 2)
limitations_rl = [
    ("Simulated Environment", "Rooms 101, 103-105 use simulated Gaussian occupancy curves, not real camera data. Curves are reasonable approximations of college schedules."),
    ("Aggressive Agent", "PPO sometimes turns everything off to save energy. Without comfort override, students could sit in the dark. Mitigated by hard-coded safety layer."),
    ("No Temperature Model", "AC decisions based on occupancy, not actual room temperature. A full BMS would also use temperature sensors."),
    ("Only 3 Device Levels", "Devices can be OFF/HALF/FULL only. Real dimmers have continuous control. Simplified to keep action space manageable."),
]
for title, desc in limitations_rl:
    add_para(title, bold=True)
    add_para(desc)

add_heading("8.3 System Limitations", 2)
add_para("Tkinter Threading: Matplotlib's TkAgg backend occasionally crashes during fast GUI updates. Wrapped in try/except — demo continues.", bold=False)
add_para("CPU-Only: All training on CPU. Mall training: ~1.5 hours, RL training: ~10 minutes. GPU would be 10-20× faster.")

doc.add_page_break()

# ============================================================
# 9. HOW TO RUN
# ============================================================
add_heading("9. How to Run", 1)

add_para("Prerequisites:", bold=True)
add_para("Python 3.10+")
add_para("pip install torch ultralytics opencv-python scipy matplotlib stable-baselines3 gymnasium tqdm")

add_para("")
add_para("Step 1: Train DL model", bold=True)
add_para("python train_mall.py")
add_para("Downloads Mall Dataset (~130MB), trains ~50 epochs, saves density_head.pth")

add_para("")
add_para("Step 2: Train RL agent", bold=True)
add_para("python rl_training/train_ppo.py --timesteps 200000")
add_para("Saves ppo_college_floor.zip")

add_para("")
add_para("Step 3: Evaluate DL model", bold=True)
add_para("python evaluate_model.py")
add_para("Generates plots in evaluation_results/")

add_para("")
add_para("Step 4: Run Demo", bold=True)
add_para("python demo.py --source aarya.mp4     (with video)")
add_para("python demo.py                        (with webcam)")
add_para("python demo.py --simulate             (simulation only)")

doc.add_page_break()

# ============================================================
# 10. FILE STRUCTURE
# ============================================================
add_heading("10. File Structure", 1)
files = [
    "demo.py — Main integration demo",
    "train_mall.py — Train density head on Mall Dataset",
    "evaluate_model.py — Evaluate trained model with plots",
    "density_head.pth — Trained density head weights",
    "ppo_college_floor.zip — Trained PPO agent",
    "dl_pipeline/model.py — YOLOv8Backbone + DensityHead",
    "dl_pipeline/preprocess.py — Density map generation",
    "rl_environment/env.py — CollegeFloorEnv (Gymnasium)",
    "rl_training/train_ppo.py — PPO training script",
    "rl_training/evaluate_rl.py — Compare PPO vs baselines",
    "rl_training/beta_sweep.py — Beta hyperparameter sweep",
    "visualization/dashboard.py — Real-time matplotlib dashboard",
    "evaluation_results/ — Model accuracy plots",
]
for f in files:
    doc.add_paragraph(f, style='List Bullet')

doc.add_page_break()

# ============================================================
# 11. REFERENCES
# ============================================================
add_heading("11. References", 1)
refs = [
    'C. C. Loy, K. Chen, S. Gong, T. Xiang. "From Semi-Supervised to Transfer Counting of Crowds." ICCV 2013. (Mall Dataset)',
    'Jocher, G., Chaurasia, A., Qiu, J. (2023). "Ultralytics YOLOv8." GitHub. (Backbone)',
    'Schulman, J., et al. "Proximal Policy Optimization Algorithms." arXiv:1707.06347, 2017. (PPO)',
    'Raffin, A., et al. "Stable-Baselines3: Reliable RL Implementations." JMLR, 2021. (SB3)',
    'Farama Foundation. "Gymnasium: A Standard API for RL." 2023. (Environment)',
    'Zhang, Y., et al. "Single-Image Crowd Counting via Multi-Column CNN." CVPR, 2016. (Density Maps)',
    'Yu, F., Koltun, V. "Multi-Scale Context Aggregation by Dilated Convolutions." ICLR, 2016. (Dilated CNN)',
]
for i, ref in enumerate(refs, 1):
    doc.add_paragraph(f"[{i}] {ref}")

# Save
output_path = Path(r"c:\Users\sheu\OneDrive\Documents\projects\Mega_Project\PROJECT_DOCUMENTATION.docx")
doc.save(str(output_path))
print(f"Document saved to: {output_path}")
